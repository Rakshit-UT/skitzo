"""
Document Processing Service
Handles PDF, DOCX parsing from blob URLs
"""

import fitz  # PyMuPDF
import requests
import tempfile
import os
from docx import Document
import logging
from typing import List, Dict, Any
from urllib.parse import urlparse
import re

logger = logging.getLogger(__name__)

class DocumentChunk:
    """Represents a document chunk with metadata"""

    def __init__(self, text: str, chunk_id: str, source_meta: Dict[str, Any]):
        self.text = text
        self.chunk_id = chunk_id
        self.source_meta = source_meta

    def to_dict(self) -> Dict[str, Any]:
        return {
            "text": self.text,
            "chunk_id": self.chunk_id,
            "source_meta": self.source_meta
        }

class DocumentProcessor:
    """Main document processing service"""

    def __init__(self):
        self.chunk_size = 1000
        self.overlap = 200

    async def process_document_from_url(self, url: str) -> List[DocumentChunk]:
        """Download and process document from URL"""
        logger.info(f"Processing document from URL: {url}")

        try:
            # Download the document
            response = requests.get(url, timeout=30)
            response.raise_for_status()

            # Determine file type
            file_extension = self._get_file_extension(url, response.headers.get('content-type', ''))

            # Save to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
                tmp_file.write(response.content)
                tmp_file_path = tmp_file.name

            try:
                # Process based on file type
                if file_extension == '.pdf':
                    chunks = self._process_pdf(tmp_file_path, url)
                elif file_extension in ['.docx', '.doc']:
                    chunks = self._process_docx(tmp_file_path, url)
                else:
                    chunks = self._process_text(tmp_file_path, url)

                logger.info(f"Successfully processed document into {len(chunks)} chunks")
                return chunks

            finally:
                os.unlink(tmp_file_path)

        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            raise Exception(f"Failed to process document: {str(e)}")

    def _get_file_extension(self, url: str, content_type: str) -> str:
        """Determine file extension from URL or content type"""
        parsed_url = urlparse(url)
        path = parsed_url.path.lower()

        if path.endswith('.pdf'):
            return '.pdf'
        elif path.endswith('.docx'):
            return '.docx'
        elif 'pdf' in content_type.lower():
            return '.pdf'
        elif 'word' in content_type.lower():
            return '.docx'

        return '.pdf'  # Default to PDF

    def _process_pdf(self, file_path: str, source_url: str) -> List[DocumentChunk]:
        """Process PDF document"""
        chunks = []

        try:
            doc = fitz.open(file_path)

            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()

                if page_text.strip():
                    page_chunks = self._create_text_chunks(
                        page_text,
                        {
                            "source_url": source_url,
                            "page": page_num + 1,
                            "document_type": "pdf",
                            "total_pages": len(doc)
                        }
                    )
                    chunks.extend(page_chunks)

            doc.close()

        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            raise Exception(f"Failed to process PDF: {str(e)}")

        return chunks

    def _process_docx(self, file_path: str, source_url: str) -> List[DocumentChunk]:
        """Process DOCX document"""
        chunks = []

        try:
            doc = Document(file_path)
            full_text = ""

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text += paragraph.text + "\n\n"

            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text += " | ".join(row_text) + "\n"

            if full_text.strip():
                chunks = self._create_text_chunks(
                    full_text,
                    {
                        "source_url": source_url,
                        "document_type": "docx"
                    }
                )

        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise Exception(f"Failed to process DOCX: {str(e)}")

        return chunks

    def _process_text(self, file_path: str, source_url: str) -> List[DocumentChunk]:
        """Process plain text document"""
        chunks = []

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()

            if text.strip():
                chunks = self._create_text_chunks(
                    text,
                    {
                        "source_url": source_url,
                        "document_type": "text"
                    }
                )

        except Exception as e:
            logger.error(f"Error processing text file: {str(e)}")
            raise Exception(f"Failed to process text file: {str(e)}")

        return chunks

    def _create_text_chunks(self, text: str, base_metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Create overlapping text chunks"""
        chunks = []

        if len(text) <= self.chunk_size:
            chunk = DocumentChunk(
                text=text,
                chunk_id=f"chunk_0",
                source_meta={**base_metadata, "chunk_index": 0}
            )
            chunks.append(chunk)
        else:
            start = 0
            chunk_index = 0

            while start < len(text):
                end = start + self.chunk_size

                if end < len(text):
                    break_point = text.rfind(' ', end - self.overlap, end)
                    if break_point > start:
                        end = break_point

                chunk_text = text[start:end].strip()

                if chunk_text:
                    chunk = DocumentChunk(
                        text=chunk_text,
                        chunk_id=f"chunk_{chunk_index}",
                        source_meta={
                            **base_metadata,
                            "chunk_index": chunk_index,
                            "start_char": start,
                            "end_char": end
                        }
                    )
                    chunks.append(chunk)
                    chunk_index += 1

                start = end - self.overlap if end < len(text) else len(text)

        return chunks
